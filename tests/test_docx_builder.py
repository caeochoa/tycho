"""Tests for the DOCX resume builder."""

from pathlib import Path

import pytest
from docx import Document

from tycho.cv.docx_builder import build_docx
from tycho.cv.module_selector import select_modules


class TestBuildDocx:
    def test_file_created(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        result = build_docx(tailored, output_path)
        assert result.exists()
        assert result.suffix == ".docx"

    def test_contains_name(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Cesar Ochoa" in all_text

    def test_contains_sections(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Professional Summary" in all_text
        assert "Technical Skills" in all_text
        assert "Work Experience" in all_text
        assert "Education" in all_text
        assert "Languages" in all_text

    def test_contains_skills(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Python" in all_text

    def test_contains_bullets(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Should contain at least one bullet from experience
        assert "RAG" in all_text or "computer vision" in all_text or "OKM" in all_text

    def test_calibri_font(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        # Check default style font
        normal_style = doc.styles["Normal"]
        assert normal_style.font.name == "Calibri"

    def test_no_images(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path)

        doc = Document(str(output_path))
        # Check no inline shapes (images)
        for para in doc.paragraphs:
            for run in para.runs:
                assert len(run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) == 0

    def test_spanish_headings(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job, language="es")
        output_path = tmp_path / "test_cv_es.docx"
        build_docx(tailored, output_path, language="es")

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Resumen Profesional" in all_text
        assert "Habilidades Técnicas" in all_text
        assert "Experiencia Laboral" in all_text

    def test_phone_selection_spain(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path, country="Spain")

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "+34" in all_text

    def test_phone_selection_uk(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "test_cv.docx"
        build_docx(tailored, output_path, country="UK")

        doc = Document(str(output_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "+44" in all_text

    def test_creates_parent_dirs(self, sample_profile, ml_job, tmp_path):
        tailored = select_modules(sample_profile, ml_job)
        output_path = tmp_path / "deep" / "nested" / "dir" / "test_cv.docx"
        result = build_docx(tailored, output_path)
        assert result.exists()
