"""Tests for the cover letter generator."""

from pathlib import Path
from unittest.mock import MagicMock

import pytest

from tycho.config import CoverLetterConfig
from tycho.cover_letter.generator import (
    _parse_response,
    generate_cover_letter,
    save_cover_letter,
)
from tycho.cv.module_selector import select_modules
from tycho.models import CoverLetter


class TestParseResponse:
    def test_structured_response(self):
        response = (
            "GREETING: Dear Hiring Team,\n"
            "PARAGRAPH: I am excited to apply for this role.\n"
            "PARAGRAPH: My experience in ML makes me a strong candidate.\n"
            "PARAGRAPH: I look forward to discussing this opportunity.\n"
            "CLOSING: Best regards,"
        )
        cl = _parse_response(response, "job-123", "en")
        assert cl.greeting == "Dear Hiring Team,"
        assert len(cl.paragraphs) == 3
        assert cl.closing == "Best regards,"
        assert cl.job_id == "job-123"
        assert cl.language == "en"

    def test_fallback_parsing(self):
        response = "This is a raw cover letter without markers."
        cl = _parse_response(response, "job-123", "en")
        # Should fallback to raw text as single paragraph
        assert len(cl.paragraphs) == 1
        assert cl.paragraphs[0] == response.strip()
        # Default greeting and closing
        assert cl.greeting == "Dear Hiring Manager,"
        assert cl.closing == "Sincerely,"

    def test_partial_markers(self):
        response = (
            "GREETING: Hello,\n"
            "Some text without paragraph marker.\n"
        )
        cl = _parse_response(response, "job-123", "en")
        assert cl.greeting == "Hello,"
        # No paragraphs parsed → fallback to full text
        assert len(cl.paragraphs) == 1

    def test_empty_response(self):
        cl = _parse_response("", "job-123", "en")
        assert len(cl.paragraphs) == 1
        assert cl.paragraphs[0] == ""


class TestCoverLetterModel:
    def test_full_text_property(self):
        cl = CoverLetter(
            job_id="job-123",
            greeting="Dear Team,",
            paragraphs=["First paragraph.", "Second paragraph."],
            closing="Best,",
        )
        text = cl.full_text
        assert "Dear Team," in text
        assert "First paragraph." in text
        assert "Second paragraph." in text
        assert "Best," in text


class TestGenerateCoverLetter:
    def test_generation_with_mock_llm(
        self, sample_profile, ml_job, mock_llm_client, cover_letter_config
    ):
        tailored = select_modules(sample_profile, ml_job)

        # Configure mock to return structured response
        mock_llm_client.invoke.return_value = (
            "GREETING: Dear Hiring Manager,\n"
            "PARAGRAPH: I am writing to express my interest.\n"
            "PARAGRAPH: My background in AI is a strong fit.\n"
            "PARAGRAPH: I look forward to your response.\n"
            "CLOSING: Sincerely,"
        )

        cl = generate_cover_letter(
            job=ml_job,
            profile=sample_profile,
            tailored=tailored,
            llm_client=mock_llm_client,
            config=cover_letter_config,
        )

        assert isinstance(cl, CoverLetter)
        assert cl.job_id == ml_job.id
        assert len(cl.paragraphs) == 3
        mock_llm_client.invoke.assert_called_once()

    def test_generation_spanish(
        self, sample_profile, ml_job, mock_llm_client, cover_letter_config
    ):
        tailored = select_modules(sample_profile, ml_job, language="es")

        mock_llm_client.invoke.return_value = (
            "GREETING: Estimado equipo,\n"
            "PARAGRAPH: Me dirijo a ustedes.\n"
            "CLOSING: Atentamente,"
        )

        cl = generate_cover_letter(
            job=ml_job,
            profile=sample_profile,
            tailored=tailored,
            llm_client=mock_llm_client,
            config=cover_letter_config,
            language="es",
        )

        assert cl.language == "es"


class TestSaveCoverLetter:
    def test_save_txt(self, tmp_path):
        cl = CoverLetter(
            job_id="job-123",
            greeting="Dear Team,",
            paragraphs=["First paragraph.", "Second paragraph."],
            closing="Best,",
        )
        output = tmp_path / "cover_letter.txt"
        result = save_cover_letter(cl, output, format="txt")
        assert result.exists()
        content = result.read_text()
        assert "Dear Team," in content
        assert "First paragraph." in content
        assert "Best," in content

    def test_save_docx(self, tmp_path):
        cl = CoverLetter(
            job_id="job-123",
            greeting="Dear Team,",
            paragraphs=["First paragraph.", "Second paragraph."],
            closing="Best,",
        )
        output = tmp_path / "cover_letter.docx"
        result = save_cover_letter(cl, output, format="docx")
        assert result.exists()
        assert result.suffix == ".docx"

        from docx import Document
        doc = Document(str(result))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Dear Team," in all_text
        assert "First paragraph." in all_text
        assert "Best," in all_text

    def test_creates_parent_dirs(self, tmp_path):
        cl = CoverLetter(
            job_id="job-123",
            paragraphs=["Test"],
        )
        output = tmp_path / "deep" / "nested" / "cl.txt"
        result = save_cover_letter(cl, output)
        assert result.exists()
