"""ATS-friendly .docx resume builder using python-docx."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from tycho.models import TailoredProfile

# Section headings by language
_HEADINGS = {
    "en": {
        "summary": "Professional Summary",
        "skills": "Technical Skills",
        "experience": "Work Experience",
        "education": "Education",
        "other": "Other Experience",
        "languages": "Languages",
    },
    "es": {
        "summary": "Resumen Profesional",
        "skills": "Habilidades Técnicas",
        "experience": "Experiencia Laboral",
        "education": "Educación",
        "other": "Otra Experiencia",
        "languages": "Idiomas",
    },
}


def build_docx(
    profile: TailoredProfile,
    output_path: str | Path,
    language: str = "en",
    country: str = "Spain",
) -> Path:
    """Build an ATS-friendly .docx resume from a tailored profile."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    headings = _HEADINGS.get(language, _HEADINGS["en"])

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header: Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile.personal.name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = "Calibri"

    # Header: Titles
    if profile.personal.titles:
        titles_text = " | ".join(profile.personal.titles)
        titles_para = doc.add_paragraph()
        titles_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titles_run = titles_para.add_run(titles_text)
        titles_run.font.size = Pt(10)
        titles_run.font.name = "Calibri"
        titles_run.font.color.rgb = RGBColor(80, 80, 80)

    # Header: Contact info
    phone = profile.personal.phone_es if country == "Spain" else profile.personal.phone_uk
    contact_parts = []
    if profile.personal.email:
        contact_parts.append(profile.personal.email)
    if phone:
        contact_parts.append(phone)
    if profile.personal.linkedin:
        contact_parts.append(profile.personal.linkedin)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.name = "Calibri"

    # Summary
    if profile.summary:
        doc.add_heading(headings["summary"], level=1)
        doc.add_paragraph(profile.summary)

    # Skills
    if profile.skills:
        doc.add_heading(headings["skills"], level=1)
        doc.add_paragraph(", ".join(profile.skills))

    # Work Experience
    if profile.experience:
        doc.add_heading(headings["experience"], level=1)
        for entry in profile.experience:
            _add_entry(doc, entry)

    # Education
    if profile.education:
        doc.add_heading(headings["education"], level=1)
        for entry in profile.education:
            _add_entry(doc, entry, show_gpa=True)

    # Other
    if profile.other:
        doc.add_heading(headings["other"], level=1)
        for entry in profile.other:
            _add_entry(doc, entry)

    # Languages
    if profile.languages:
        doc.add_heading(headings["languages"], level=1)
        lang_items = [f"{lang.language}: {lang.level}" for lang in profile.languages]
        doc.add_paragraph(", ".join(lang_items))

    doc.save(str(output_path))
    return output_path


def _add_entry(doc: Document, entry, show_gpa: bool = False) -> None:
    """Add a resume entry (experience, education, or other) to the document."""
    # Title line: Title — Organization | Dates
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f"{entry.title}")
    title_run.bold = True
    title_run.font.size = Pt(10)
    title_run.font.name = "Calibri"

    org_run = title_para.add_run(f" — {entry.organization}")
    org_run.font.size = Pt(10)
    org_run.font.name = "Calibri"

    if entry.dates:
        dates_run = title_para.add_run(f" | {entry.dates}")
        dates_run.font.size = Pt(10)
        dates_run.font.name = "Calibri"
        dates_run.font.color.rgb = RGBColor(100, 100, 100)

    if entry.location:
        loc_run = title_para.add_run(f" | {entry.location}")
        loc_run.font.size = Pt(9)
        loc_run.font.name = "Calibri"
        loc_run.font.color.rgb = RGBColor(100, 100, 100)

    # GPA if applicable
    if show_gpa and getattr(entry, "gpa", None):
        gpa_para = doc.add_paragraph()
        gpa_run = gpa_para.add_run(f"GPA: {entry.gpa}")
        gpa_run.font.size = Pt(9)
        gpa_run.font.name = "Calibri"
        gpa_run.italic = True

    # Note if present
    if entry.note:
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(entry.note)
        note_run.font.size = Pt(9)
        note_run.font.name = "Calibri"
        note_run.italic = True

    # Bullets
    for bullet in entry.bullets:
        bullet_para = doc.add_paragraph(bullet.text, style="List Bullet")
        for run in bullet_para.runs:
            run.font.size = Pt(10)
            run.font.name = "Calibri"
