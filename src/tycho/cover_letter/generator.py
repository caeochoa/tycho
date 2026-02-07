"""LLM-based cover letter generation."""

from __future__ import annotations

import logging
from pathlib import Path

from tycho.config import CoverLetterConfig
from tycho.models import CoverLetter, Job, Profile, TailoredProfile

logger = logging.getLogger(__name__)


def generate_cover_letter(
    job: Job,
    profile: Profile,
    tailored: TailoredProfile,
    llm_client,
    config: CoverLetterConfig,
    language: str = "en",
) -> CoverLetter:
    """Generate a cover letter using the LLM.

    Returns a structured CoverLetter model.
    """
    # Build the prompt
    system_prompt = _build_system_prompt(config, language)
    human_prompt = _build_human_prompt(job, profile, tailored, language)

    full_prompt = f"{system_prompt}\n\n{human_prompt}"
    response = llm_client.invoke(full_prompt)

    # Parse structured response
    cover_letter = _parse_response(response, job.id, language)
    return cover_letter


def save_cover_letter(
    cover_letter: CoverLetter,
    output_path: str | Path,
    format: str = "txt",
) -> Path:
    """Save a cover letter to file.

    Supports 'txt' and 'docx' formats.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if format == "docx":
        return _save_docx(cover_letter, output_path)

    # Default: plain text
    output_path.write_text(cover_letter.full_text, encoding="utf-8")
    return output_path


def _build_system_prompt(config: CoverLetterConfig, language: str) -> str:
    """Build the system prompt for cover letter generation."""
    lang_instruction = "Write in Spanish." if language == "es" else "Write in English."

    return (
        f"You are a professional cover letter writer. {lang_instruction}\n"
        f"Tone: {config.tone}.\n"
        f"Structure the letter with exactly {config.max_paragraphs} body paragraphs.\n\n"
        "Format your response using these markers:\n"
        "GREETING: <the greeting line>\n"
        "PARAGRAPH: <first body paragraph>\n"
        "PARAGRAPH: <second body paragraph>\n"
        "PARAGRAPH: <third body paragraph>\n"
        "CLOSING: <the closing line>"
    )


def _build_human_prompt(
    job: Job, profile: Profile, tailored: TailoredProfile, language: str
) -> str:
    """Build the human prompt with job and candidate context."""
    # Top 3 experience entries
    top_experiences = []
    for entry in tailored.experience[:3]:
        bullets_text = "; ".join(b.text for b in entry.bullets[:2])
        top_experiences.append(f"- {entry.title} at {entry.organization}: {bullets_text}")

    experience_text = "\n".join(top_experiences) if top_experiences else "N/A"
    skills_text = ", ".join(tailored.skills[:10]) if tailored.skills else "N/A"

    return (
        f"Write a cover letter for the following job application:\n\n"
        f"Job Title: {job.title}\n"
        f"Company: {job.company}\n"
        f"Job Description:\n{job.description[:2000]}\n\n"
        f"Candidate Name: {profile.personal.name}\n"
        f"Summary: {tailored.summary}\n"
        f"Top Skills: {skills_text}\n"
        f"Key Experience:\n{experience_text}"
    )


def _parse_response(response: str, job_id: str, language: str) -> CoverLetter:
    """Parse the LLM response into a structured CoverLetter.

    Falls back to raw text as a single paragraph if parsing fails.
    """
    greeting = "Dear Hiring Manager,"
    paragraphs = []
    closing = "Sincerely,"

    lines = response.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.upper().startswith("GREETING:"):
            greeting = line[len("GREETING:"):].strip()
        elif line.upper().startswith("PARAGRAPH:"):
            paragraphs.append(line[len("PARAGRAPH:"):].strip())
        elif line.upper().startswith("CLOSING:"):
            closing = line[len("CLOSING:"):].strip()

    # Fallback: if no paragraphs parsed, use raw text
    if not paragraphs:
        paragraphs = [response.strip()]

    return CoverLetter(
        job_id=job_id,
        greeting=greeting,
        paragraphs=paragraphs,
        closing=closing,
        language=language,
    )


def _save_docx(cover_letter: CoverLetter, output_path: Path) -> Path:
    """Save cover letter as .docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Greeting
    doc.add_paragraph(cover_letter.greeting)

    # Paragraphs
    for para in cover_letter.paragraphs:
        doc.add_paragraph(para)

    # Closing
    doc.add_paragraph(cover_letter.closing)

    doc.save(str(output_path))
    return output_path
